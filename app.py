import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from dotenv import load_dotenv
import pandas as pd
from PIL import Image
import pytesseract
import io
import docx
import logging
from PIL import UnidentifiedImageError, ImageFile
ImageFile.LOAD_TRUNCATED_IMAGES = True  # Handle truncated images


# Configure logging
logging.basicConfig(level=logging.INFO, filename='app.log', filemode='a',
                    format='%(asctime)s - %(levelname)s - %(message)s')

load_dotenv()
google_api_key = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=google_api_key)

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Ensure the faiss_index folder exists
os.makedirs("faiss_index", exist_ok=True)

# Create the embedding instance once
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=google_api_key)


def get_text_and_metadata(file_docs, doc_type):
    text = ""
    metadata = []
    images = []

    if doc_type == "pdf":
        for file in file_docs:
            pdf_reader = PdfReader(file)
            doc_metadata = pdf_reader.metadata
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                    metadata.append(doc_metadata)
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            try:
                                img_data = xObject[obj]._data
                                images.append(Image.open(io.BytesIO(img_data)))
                            except UnidentifiedImageError:
                                print("Skipped unidentified image in PDF")
    
    elif doc_type == "docx":
        for file in file_docs:
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
                metadata.append({"docx_length": len(doc.paragraphs)})
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
                    metadata.append({"docx_length": len(doc.paragraphs)})
            for rel in doc.part.rels:
                if "image" in doc.part.rels[rel].target_ref:
                    try:
                        img = Image.open(io.BytesIO(doc.part.rels[rel].target_part.blob))
                        images.append(img)
                    except UnidentifiedImageError:
                        print("Skipped unidentified image in DOCX")
    
    elif doc_type == "xlsx":
        for file in file_docs:
            xls = pd.ExcelFile(file)
            sheet_names = xls.sheet_names
            for sheet_name in sheet_names:
                df = pd.read_excel(file, sheet_name=sheet_name)
                numerical_cols = df.select_dtypes(include='number').columns.tolist()
                categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
                text += df.to_string()
                metadata.append({"sheet_name": sheet_name, "shape": df.shape, "numerical_cols": numerical_cols, "categorical_cols": categorical_cols})
    
    elif doc_type == "csv":
        for file in file_docs:
            df = pd.read_csv(file)
            numerical_cols = df.select_dtypes(include='number').columns.tolist()
            categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
            text += df.to_string()
            metadata.append({"shape": df.shape, "columns": df.columns.tolist(), "numerical_cols": numerical_cols, "categorical_cols": categorical_cols})
    
    elif doc_type == "image":
        for file in file_docs:
            try:
                img = Image.open(file)
                img_text = pytesseract.image_to_string(img)
                if img_text:
                    text += img_text
                    metadata.append({"format": img.format, "size": img.size, "mode": img.mode})
            except UnidentifiedImageError:
                print(f"Skipped unidentified image: {file.name}")
    
    return text, metadata, images

def process_images(images):
    descriptions = []
    for img in images:
        try:
            text = pytesseract.image_to_string(img)
            description = f"Image with dimensions {img.size} and mode {img.mode}. Text extracted: {text}"
            descriptions.append({"text": description, "metadata": {"dimensions": img.size, "mode": img.mode}})
        except Exception as e:
            print(f"Error processing image: {e}")
    return descriptions


def get_text_chunks_and_metadata(text, metadata):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    text_chunks = text_splitter.split_text(text)
    chunk_metadata = []

    for i in range(len(text_chunks)):
        chunk_metadata.append(metadata[i % len(metadata)])  # Cycle metadata if necessary
    
    return text_chunks, chunk_metadata

def create_and_save_vector_store(text_chunks, metadata):
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings, metadatas=metadata)
    vector_store.save_local("faiss_index")
    logging.info("Vector store created and saved successfully.")
    return vector_store

def load_vector_store():
    vector_store = FAISS.load_local("faiss_index", embeddings)
    logging.info("Vector store loaded successfully.")
    return vector_store

def get_advanced_prompt():
    return '''
    You are a real estate expert providing detailed and accurate information. Answer the questions based on the provided context which includes property details, legal checklists, sales policies, and inventory data. Be as comprehensive and precise as possible. If the information is not available, state "The information is not available in the provided context".

    Context:
    {context}

    Question: 
    {question}

    Answer:
    '''

def get_conversational_chain(memory):
    prompt_template = get_advanced_prompt()
    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3, google_api_key=google_api_key)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt, memory=memory)

def user_input(user_question, vector_store, memory):
    context_documents = vector_store.similarity_search(user_question)
    context = [doc.page_content for doc in context_documents]
    
    combined_input = {
        "input_documents": context_documents,  # Ensure input_documents is provided
        "context": context, 
        "question": user_question
    }
    
    chain = get_conversational_chain(memory)
    response = chain.invoke(combined_input)
    
    # Extract relevant parts for display
    display_response = {
        "question": user_question,
        "chat_history": memory.chat_memory,  # Correctly accessing chat history
        "output_text": response['output_text']
    }
    
    return display_response

def main():
    st.title("Document Chat with Multiple Formats")

    if "vector_store" not in st.session_state:
        st.session_state.vector_store = None
    if "metadata" not in st.session_state:
        st.session_state.metadata = []
    if "memory" not in st.session_state:
        st.session_state.memory = ConversationBufferMemory(memory_key="chat_history", input_key="question")

    pdf_docs = st.file_uploader("Upload PDF Files", type=["pdf"], accept_multiple_files=True)
    docx_docs = st.file_uploader("Upload DOCX Files", type=["docx"], accept_multiple_files=True)
    xlsx_docs = st.file_uploader("Upload Excel Files", type=["xlsx"], accept_multiple_files=True)
    csv_docs = st.file_uploader("Upload CSV Files", type=["csv"], accept_multiple_files=True)
    image_docs = st.file_uploader("Upload Image Files", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

    if st.button("Process Documents"):
        text = ""
        metadata = []
        images = []

        if pdf_docs:
            pdf_text, pdf_metadata, pdf_images = get_text_and_metadata(pdf_docs, "pdf")
            text += pdf_text
            metadata.extend(pdf_metadata)
            images.extend(pdf_images)
        if docx_docs:
            docx_text, docx_metadata, docx_images = get_text_and_metadata(docx_docs, "docx")
            text += docx_text
            metadata.extend(docx_metadata)
            images.extend(docx_images)
        if xlsx_docs:
            xlsx_text, xlsx_metadata, _ = get_text_and_metadata(xlsx_docs, "xlsx")
            text += xlsx_text
            metadata.extend(xlsx_metadata)
        if csv_docs:
            csv_text, csv_metadata, _ = get_text_and_metadata(csv_docs, "csv")
            text += csv_text
            metadata.extend(csv_metadata)
        if image_docs:
            image_text, image_metadata, _ = get_text_and_metadata(image_docs, "image")
            text += image_text
            metadata.extend(image_metadata)

        if text:
            text_chunks, chunk_metadata = get_text_chunks_and_metadata(text, metadata)
            image_descriptions = process_images(images)
            for desc in image_descriptions:
                text_chunks.append(desc["text"])
                chunk_metadata.append(desc["metadata"])

            st.write(f"Number of text chunks: {len(text_chunks)}")
            st.write(f"Number of metadata items: {len(chunk_metadata)}")

            if len(text_chunks) != len(chunk_metadata):
                st.error("The number of text chunks and metadata items do not match!")
            else:
                vector_store = create_and_save_vector_store(text_chunks, chunk_metadata)
                st.session_state.vector_store = vector_store
                st.session_state.metadata = chunk_metadata
                st.success("Documents processed and vector store created successfully!")

    with st.form("document_question_form"):
        user_question = st.text_input("Ask a question about the documents")
        submit_button = st.form_submit_button("Ask Question")

        if submit_button:
            if st.session_state.vector_store is None:
                st.session_state.vector_store = load_vector_store()
                st.success("Vector store loaded successfully!")

            if st.session_state.vector_store:
                answer = user_input(user_question, st.session_state.vector_store, st.session_state.memory)
                # Display only the relevant sections
                st.write("**Question:**")
                st.write(answer['question'])
                st.write("**Chat History:**")
                st.write(answer['chat_history'])
                st.write("**Answer:**")
                st.write(answer['output_text'])

if __name__ == "__main__":
    main()
